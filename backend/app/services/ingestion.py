"""Document ingestion: save upload and extract text from TXT, PDF, DOCX."""
import re
import uuid
import chardet
from pathlib import Path
from typing import Tuple
from fastapi import UploadFile, HTTPException

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

MAX_FILE_BYTES = 500 * 1024 * 1024          # 500 MB hard limit
ALLOWED_EXTENSIONS = {"txt", "pdf", "docx"}


def _safe_filename(filename: str) -> str:
    name = re.sub(r"[^\w\-.]", "_", filename)
    return f"{uuid.uuid4().hex}_{name}"


async def save_upload(file: UploadFile) -> Tuple[Path, str]:
    """Validate, save the uploaded file to disk and return (path, safe_name)."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided.")

    ext = Path(file.filename).suffix.lstrip(".").lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"'.{ext}' is not supported. Please upload TXT, PDF, or DOCX.",
        )

    content = await file.read()

    if len(content) == 0:
        raise HTTPException(status_code=400, detail="The uploaded file is empty.")

    if len(content) > MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds the 500 MB limit.")

    safe_name = _safe_filename(file.filename)
    save_path = UPLOAD_DIR / safe_name
    save_path.write_bytes(content)
    return save_path, safe_name


def extract_text_from_file(file_path: Path, ext: str) -> str:
    """Extract plain text. Tries multiple strategies per format before giving up."""
    if ext == "txt":
        return _extract_txt(file_path)
    elif ext == "pdf":
        return _extract_pdf(file_path)
    elif ext == "docx":
        return _extract_docx(file_path)
    raise HTTPException(status_code=400, detail=f"Unsupported extension: {ext}")


# ─── TXT ──────────────────────────────────────────────────────────────────────

def _extract_txt(path: Path) -> str:
    raw = path.read_bytes()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding") or "utf-8"
    try:
        return raw.decode(encoding, errors="replace")
    except Exception:
        return raw.decode("utf-8", errors="replace")


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> str:
    text = ""

    # Strategy 1: PyMuPDF (best quality, handles most PDFs)
    try:
        import fitz
        doc = fitz.open(str(path))
        pages = [page.get_text("text") for page in doc]
        doc.close()
        text = "\n\n".join(p for p in pages if p.strip())
        if len(text.strip()) > 50:
            return text
    except Exception:
        pass

    # Strategy 2: PyPDF2 fallback
    try:
        from PyPDF2 import PdfReader
        import io
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        text = "\n\n".join(pages)
        if len(text.strip()) > 50:
            return text
    except Exception:
        pass

    # Strategy 3: pdfminer (most compatible with complex layouts)
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(str(path))
        if len(text.strip()) > 50:
            return text
    except Exception:
        pass

    # If we got something, return it even if short
    if text.strip():
        return text

    raise HTTPException(
        status_code=422,
        detail=(
            "Could not extract text from this PDF. "
            "It may be image-based (scanned). "
            "Try converting it to text first, or use a PDF with selectable text."
        ),
    )


def extract_structured_pages(file_path: Path, ext: str) -> dict:
    """Extract page-aware content. Never fabricates page numbers.

    Returns:
      {
        "pages": [{"page": int, "raw_text": str, "blocks": [...]}],
        "raw_text": str,
        "error": optional str,
      }
    """
    if ext == "pdf":
        return _extract_pdf_pages(file_path)
    if ext == "txt":
        text = _extract_txt(file_path)
        return {
            "pages": [{"page": 1, "raw_text": text, "blocks": _blocks_from_plain(text)}],
            "raw_text": text,
        }
    if ext == "docx":
        text = _extract_docx(file_path)
        return {
            "pages": [{"page": 1, "raw_text": text, "blocks": _blocks_from_plain(text)}],
            "raw_text": text,
        }
    raise HTTPException(status_code=400, detail=f"Unsupported extension: {ext}")


def _blocks_from_plain(text: str) -> list:
    blocks = []
    for para in text.split("\n"):
        stripped = para.strip()
        if not stripped:
            continue
        kind = "heading" if _looks_like_heading(stripped) else "paragraph"
        blocks.append({"type": kind, "text": stripped})
    return blocks or [{"type": "paragraph", "text": text.strip()}] if text.strip() else []


def _looks_like_heading(line: str) -> bool:
    if len(line) > 90:
        return False
    if line.isupper() and len(line.split()) <= 12:
        return True
    if line[:1].isdigit() and ("chapter" in line.lower() or len(line.split()) <= 10):
        return True
    lowered = line.lower()
    if lowered.startswith(("chapter ", "unit ", "topic ", "section ", "module ")):
        return True
    return False


def _extract_pdf_pages(path: Path) -> dict:
    pages = []
    errors = []

    try:
        import fitz
        doc = fitz.open(str(path))
        for i, page in enumerate(doc):
            page_no = i + 1
            raw = page.get_text("text") or ""
            blocks = _blocks_from_pymupdf(page)
            pages.append({"page": page_no, "raw_text": raw, "blocks": blocks})
        doc.close()
        raw_text = "\n\n".join(p["raw_text"] for p in pages)
        return {"pages": pages, "raw_text": raw_text}
    except Exception as exc:
        errors.append(f"PyMuPDF: {exc}")

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        for i, page in enumerate(reader.pages):
            raw = page.extract_text() or ""
            pages.append({
                "page": i + 1,
                "raw_text": raw,
                "blocks": _blocks_from_plain(raw),
            })
        raw_text = "\n\n".join(p["raw_text"] for p in pages)
        return {"pages": pages, "raw_text": raw_text}
    except Exception as exc:
        errors.append(f"PyPDF2: {exc}")

    return {
        "pages": [],
        "raw_text": "",
        "error": "; ".join(errors) if errors else "PDF extraction failed",
    }


def _blocks_from_pymupdf(page) -> list:
    blocks = []
    try:
        data = page.get_text("dict")
        sizes = []
        for b in data.get("blocks", []):
            for line in b.get("lines", []):
                for span in line.get("spans", []):
                    if span.get("size"):
                        sizes.append(span["size"])
        median = sorted(sizes)[len(sizes) // 2] if sizes else 11
        for b in data.get("blocks", []):
            texts = []
            max_size = 0
            flags = 0
            for line in b.get("lines", []):
                line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                if line_text.strip():
                    texts.append(line_text.strip())
                for span in line.get("spans", []):
                    max_size = max(max_size, span.get("size") or 0)
                    flags |= span.get("flags") or 0
            text = " ".join(texts).strip()
            if not text:
                continue
            is_heading = max_size >= median + 1.5 or (flags & 16 and len(text) < 90)
            if not is_heading:
                is_heading = _looks_like_heading(text)
            blocks.append({"type": "heading" if is_heading else "paragraph", "text": text})
    except Exception:
        raw = page.get_text("text") or ""
        return _blocks_from_plain(raw)
    return blocks


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))

        parts = []
        # Main paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n\n".join(parts)
        if len(text.strip()) > 0:
            return text

        raise HTTPException(status_code=422, detail="No text content found in this DOCX file.")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"DOCX extraction failed: {str(e)}")
